import streamlit as st
from pathlib import Path
import pandas as pd
import base64
from docx import Document
from together import Together
from dotenv import load_dotenv
import os
import fitz
import io

def get_file_extension(file_name):
    return Path(file_name).suffix.lower()

def file_handler(upload_file):
    ext = get_file_extension(upload_file.name)

    if ext == ".csv":
        df = pd.read_csv(upload_file)
        st.dataframe(df)
        return df.to_string()

    elif ext == ".txt":
        text = upload_file.read().decode("utf-8")
        # lines = text.splitlines()
        # preview = "\n".join(lines[:5]) if len(lines) >= 5 else text
        st.text_area("preview", text)
        return text

    elif ext == ".pdf":
        file_bytes = upload_file.read()

        base64_pdf = base64.b64encode(file_bytes).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

        doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += page.get_text()

        # st.text_area("Extracted Text", full_text[:2000], height=300)
        return full_text

    elif ext == ".xlsx":
        xls = pd.ExcelFile(upload_file)
        sheet = st.selectbox("Select sheet", xls.sheet_names)
        df = xls.parse(sheet)
        st.dataframe(df)
        return df.to_string()

    elif ext == ".docx":
        doc = Document(upload_file)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        st.text_area("File Content", full_text, height=200)
        return full_text

    else:
        st.error("Unsupported file type.")
        return ""

def generate_response(chat_history, file_data, query):
    load_dotenv()
    client = Together(api_key="566c1ecd7169d2bdf5fcb9d71fe63148cb81deeec96fc7bd9e9ff4d4f9106d6d")
    messages = [
        {"role": "assistant", "content": "You are a helpful assistant. Use the file content to answer questions."},
        {"role": "user", "content": f"The following file content should be used as context:\n{file_data}"}
    ] + chat_history + [{"role":"user", "content":query}]
    # for role, content in chat_history:
    #     messages.append({
    #         "role": role,
    #         "content": content
    #     })
    # messages.append({"role": "user", "content": query})
    response = client.chat.completions.create(
        model="meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8",
        messages=messages
    )
    response_dict = response.model_dump()
    answer = response_dict["choices"][0]["message"]["content"]
    return answer

def doc_download(text):
    buffer = io.BytesIO()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(page_title="AskMyFile")
    st.title("🗂️ AskMyFile")

    if "chat" not in st.session_state:
        st.session_state.chat = [{"role":"assistant", "content":"How can I help you?"}]
    if "file_data" not in st.session_state:
        st.session_state.file_data = ""
    if "user_input" not in st.session_state:
        st.session_state.user_input = ""

    file = st.file_uploader("Upload file", type=["docx", "csv", "xlsx", "pdf", "txt"])
    if file:
        data = file_handler(file)
        if data:
            st.session_state.file_data = data
            st.success("File uploaded successfully!")

    st.markdown("---")
    st.subheader("Ask a question about the uploaded file:")

    # for role, msg in st.session_state.chat:
    #     if role == "user":
    #         st.markdown(f"**You:** {msg}")
    #     else:
    #         st.markdown(f"**Bot:** {msg}")


    # query = st.text_input("Your question", key="chat_input", value=st.session_state.user_input)
    if prompt := st.chat_input("Ask a question about the uploaded file:"):
        if not st.session_state.file_data:
            st.warning("📄Please upload a file.")
        else:
            with st.spinner("Thinking..."):
                st.session_state.chat.append({"role":"user", "content":prompt})
                st.chat_message("user").write(prompt)
                answer = generate_response(st.session_state.chat, st.session_state.file_data, prompt)
                st.session_state.chat.append({"role":"assistant", "content": answer})
                st.chat_message("assistant").write(answer)

                if any(word in prompt.lower() for word in ["quiz", "generate questions", "mcq", "questions"]):
                    doc_buffer = doc_download(answer)
                    st.download_button(
                        label = "📥 Download questions/Quiz",
                        data = doc_buffer,
                        file_name = "generated questions/quiz.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        # st.session_state.user_input = ""
        # st.rerun()

    # for role, msg in st.session_state.chat:
    #     if role == "user":
    #         st.markdown(f"**You:** {msg}")
    #     else:
    #         st.markdown(f"**Bot:** {msg}")

    for msg in st.session_state.chat:
        st.chat_message(msg["role"]).write(msg["content"])

if __name__ == "__main__":
    main()
